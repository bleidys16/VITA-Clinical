from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT

doc = Document()
s = doc.styles['Normal']
s.font.name = 'Calibri'
s.font.size = Pt(11)

doc.add_heading('2. Acceso al Sistema', level=2)

# 2.1 Credenciales
doc.add_heading('2.1 Credenciales de Acceso', level=3)

p = doc.add_paragraph()
p.add_run('El sistema está desplegado en la nube y se accede desde cualquier navegador web. También puede ejecutarse localmente clonando el repositorio.').font.size = Pt(11)

doc.add_paragraph()
doc.add_heading('Opción A — Acceso Web (desplegado)', level=4)

p = doc.add_paragraph()
p.add_run('Ingrese desde el navegador a la URL proporcionada por el equipo de VITA Clinical. Use las siguientes credenciales de prueba:').font.size = Pt(11)

tabla = doc.add_table(rows=4, cols=3)
tabla.style = 'Light Grid Accent 1'
tabla.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, (a, b, c) in enumerate([
    ('Rol', 'Usuario', 'Contraseña'),
    ('Administrador', 'admin@vita.com', 'admin123'),
    ('Médico', 'medico@vita.com', 'medico123'),
    ('Analista', 'analista@vita.com', 'analista123'),
]):
    for j, txt in enumerate([a, b, c]):
        run = tabla.cell(i, j).paragraphs[0].add_run(txt)
        run.font.size = Pt(10)
        if i == 0:
            run.bold = True

p = doc.add_paragraph()
r = p.add_run('(Captura de pantalla: pantalla de inicio de sesión con campos de email y contraseña)')
r.italic = True
r.font.size = Pt(10)
r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

doc.add_paragraph()

doc.add_heading('Opción B — Ejecución local (git clone)', level=4)

p = doc.add_paragraph()
p.add_run('Requisitos previos: ').bold = True
p.add_run('Python 3.12+, Git, PostgreSQL 16').font.size = Pt(11)

doc.add_paragraph()

p = doc.add_paragraph()
p.add_run('Paso 1 — Clonar el repositorio:').bold = True
p.add_run('\n\n').font.size = Pt(11)

# Code block style
codigo = doc.add_paragraph()
codigo.paragraph_format.left_indent = Pt(20)
r = codigo.add_run('git clone https://github.com/bleidys16/VITA-Clinical.git\ncd VITA-Clinical/backend')
r.font.name = 'Consolas'
r.font.size = Pt(9.5)

doc.add_paragraph()

p = doc.add_paragraph()
p.add_run('Paso 2 — Crear entorno virtual e instalar dependencias:').bold = True

codigo2 = doc.add_paragraph()
codigo2.paragraph_format.left_indent = Pt(20)
r2 = codigo2.add_run('python -m venv venv\nvenv\\Scripts\\activate\npip install -r requirements.txt')
r2.font.name = 'Consolas'
r2.font.size = Pt(9.5)

doc.add_paragraph()

p = doc.add_paragraph()
p.add_run('Paso 3 — Configurar variables de entorno:').bold = True

codigo3 = doc.add_paragraph()
codigo3.paragraph_format.left_indent = Pt(20)
r3 = codigo3.add_run('cp .env.example .env\n# Editar .env con credenciales de BD')
r3.font.name = 'Consolas'
r3.font.size = Pt(9.5)

doc.add_paragraph()

p = doc.add_paragraph()
p.add_run('Paso 4 — Migrar base de datos y crear usuarios:').bold = True

codigo4 = doc.add_paragraph()
codigo4.paragraph_format.left_indent = Pt(20)
r4 = codigo4.add_run('python manage.py migrate\npython manage.py crear_usuarios_base')
r4.font.name = 'Consolas'
r4.font.size = Pt(9.5)

doc.add_paragraph()

p = doc.add_paragraph()
p.add_run('Paso 5 — Iniciar servidor:').bold = True

codigo5 = doc.add_paragraph()
codigo5.paragraph_format.left_indent = Pt(20)
r5 = codigo5.add_run('python manage.py runserver')
r5.font.name = 'Consolas'
r5.font.size = Pt(9.5)

doc.add_paragraph()

p = doc.add_paragraph()
p.add_run('Abrir en el navegador: ').font.size = Pt(11)
r = p.add_run('http://127.0.0.1:8000')
r.font.name = 'Consolas'
r.font.size = Pt(11)

doc.add_paragraph()

p = doc.add_paragraph()
p.add_run('Roles por defecto:').bold = True

tabla2 = doc.add_table(rows=4, cols=3)
tabla2.style = 'Light Grid Accent 1'
tabla2.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, (a, b, c) in enumerate([
    ('Rol', 'Usuario', 'Contraseña'),
    ('Administrador', 'admin@vita.com', 'admin123'),
    ('Médico', 'medico@vita.com', 'medico123'),
    ('Analista', 'analista@vita.com', 'analista123'),
]):
    for j, txt in enumerate([a, b, c]):
        run = tabla2.cell(i, j).paragraphs[0].add_run(txt)
        run.font.size = Pt(10)
        if i == 0:
            run.bold = True

p = doc.add_paragraph()
r = p.add_run('(Captura de pantalla: terminal mostrando el servidor corriendo en localhost)')
r.italic = True
r.font.size = Pt(10)
r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

doc.add_paragraph()

# 2.2 Sidebar
doc.add_heading('2.2 Navegación (Sidebar)', level=3)
p = doc.add_paragraph()
p.add_run('El menú lateral se adapta según el rol del usuario:').font.size = Pt(11)

tabla3 = doc.add_table(rows=6, cols=4)
tabla3.style = 'Light Grid Accent 1'
tabla3.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, (a, b, c, d) in enumerate([
    ('Sección', 'Admin', 'Médico', 'Analista'),
    ('Dashboard', '✅', '✅', '✅'),
    ('Analítica', '✅', '✅', '✅'),
    ('Pacientes', '✅', '✅', '✅'),
    ('ETL / ML', '✅', '❌', '✅'),
    ('Usuarios', '✅', '❌', '❌'),
]):
    for j, txt in enumerate([a, b, c, d]):
        run = tabla3.cell(i, j).paragraphs[0].add_run(txt)
        run.font.size = Pt(10)
        if i == 0:
            run.bold = True

p = doc.add_paragraph()
r = p.add_run('(Captura de pantalla: sidebar según rol — ADMIN viendo todas las opciones, MEDICO con opciones limitadas)')
r.italic = True
r.font.size = Pt(10)
r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

doc.save(r'C:\Users\Bleidys Larios\CLIA\Seccion2_Acceso.docx')
print("OK")
