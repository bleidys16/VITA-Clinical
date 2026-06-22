from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT

doc = Document()
s = doc.styles['Normal']
s.font.name = 'Calibri'
s.font.size = Pt(11)

doc.add_heading('5. Analítica Descriptiva', level=2)

p = doc.add_paragraph()
p.add_run('Esta sección está disponible desde el sidebar mediante el enlace ').font.size = Pt(11)
r = p.add_run('Analítica')
r.bold = True
r.font.size = Pt(11)
p.add_run(' (visible para ADMIN y ANALISTA).').font.size = Pt(11)

doc.add_paragraph()

# Bloque 1
doc.add_heading('5.1 Prevalencia Patológica', level=3)
p = doc.add_paragraph()
p.add_run('Muestra ').font.size = Pt(11)
r = p.add_run('7 indicadores')
r.bold = True
r.font.size = Pt(11)
p.add_run(' cliqueables con barra de progreso y porcentaje:').font.size = Pt(11)

t1 = doc.add_table(rows=8, cols=3)
t1.style = 'Light Grid Accent 1'
t1.alignment = WD_TABLE_ALIGNMENT.CENTER
d1 = [
    ('Indicador', 'Icono', 'Descripción'),
    ('Hipertensión Arterial', 'fa-heart-pulse', 'Presión sistólica elevada'),
    ('Diabetes Mellitus', 'fa-droplet', 'Glucosa elevada'),
    ('Tabaquismo Activo', 'fa-smoking', 'Pacientes fumadores'),
    ('Obesidad', 'fa-weight-scale', 'IMC ≥ 30'),
    ('Antecedentes Familiares', 'fa-dna', 'Con antecedentes'),
    ('Consumo de Alcohol', 'fa-wine-bottle', 'Bebedores activos'),
    ('Saturación Baja (<85%)', 'fa-lungs', 'Oxígeno bajo'),
]
for i, (a, b, c) in enumerate(d1):
    for j, txt in enumerate([a, b, c]):
        run = t1.cell(i, j).paragraphs[0].add_run(txt)
        run.font.size = Pt(9.5)
        if i == 0: run.bold = True

p2 = doc.add_paragraph()
p2.add_run('(Captura de pantalla: sección Prevalencia Patológica con las 7 tarjetas visibles)').italic = True
p2.runs[0].font.size = Pt(10)
p2.runs[0].font.color.rgb = RGBColor(0x66, 0x66, 0x66)

doc.add_paragraph()

# Bloque 2
doc.add_heading('5.2 Alertas Clínicas', level=3)
p = doc.add_paragraph()
p.add_run('3 alertas críticas con el número de pacientes afectados:').font.size = Pt(11)

t2 = doc.add_table(rows=4, cols=2)
t2.style = 'Light Grid Accent 1'
t2.alignment = WD_TABLE_ALIGNMENT.CENTER
d2 = [
    ('Alerta', 'Criterio'),
    ('Presión Sistólica > 180 mmHg', 'Pacientes en crisis hipertensiva'),
    ('Glucosa > 300 mg/dL', 'Pacientes con hiperglucemia severa'),
    ('Sat. Oxígeno < 85%', 'Pacientes con hipoxemia'),
]
for i, (a, b) in enumerate(d2):
    for j, txt in enumerate([a, b]):
        run = t2.cell(i, j).paragraphs[0].add_run(txt)
        run.font.size = Pt(9.5)
        if i == 0: run.bold = True

p3 = doc.add_paragraph()
p3.add_run('(Captura de pantalla: tarjetas de Alertas Clínicas)').italic = True
p3.runs[0].font.size = Pt(10)
p3.runs[0].font.color.rgb = RGBColor(0x66, 0x66, 0x66)

doc.add_paragraph()

# Bloque 3
doc.add_heading('5.3 Segmentación', level=3)
p = doc.add_paragraph()
p.add_run('Dos tablas con barra de progreso:').font.size = Pt(11)

t3 = doc.add_table(rows=3, cols=2)
t3.style = 'Light Grid Accent 1'
t3.alignment = WD_TABLE_ALIGNMENT.CENTER
d3 = [
    ('Segmentación', 'Categorías'),
    ('Por Nivel de Riesgo', 'Bajo · Medio · Alto · Crítico'),
    ('Por Grupo de Edad', '<30 · 30-49 · 50-69 · 70+'),
]
for i, (a, b) in enumerate(d3):
    for j, txt in enumerate([a, b]):
        run = t3.cell(i, j).paragraphs[0].add_run(txt)
        run.font.size = Pt(9.5)
        if i == 0: run.bold = True

p4 = doc.add_paragraph()
p4.add_run('(Captura de pantalla: tablas de segmentación por riesgo y edad)').italic = True
p4.runs[0].font.size = Pt(10)
p4.runs[0].font.color.rgb = RGBColor(0x66, 0x66, 0x66)

doc.add_paragraph()

# Bloque 4
doc.add_heading('5.4 Matriz Descriptiva', level=3)
p = doc.add_paragraph()
p.add_run('Tabla estadística con 9 métricas (Count, Media, Mediana, Moda, Desv. Estándar, Mín, Máx, P25, P75) para cada variable biomédica continua (glucosa, colesterol, IMC, presión, frecuencia cardíaca, saturación, temperatura, etc.).').font.size = Pt(11)

p5 = doc.add_paragraph()
p5.add_run('(Captura de pantalla: Matriz Descriptiva de Variables Biomédicas Continuas)').italic = True
p5.runs[0].font.size = Pt(10)
p5.runs[0].font.color.rgb = RGBColor(0x66, 0x66, 0x66)

doc.add_paragraph()

# Bloque 5
doc.add_heading('5.5 Distribuciones Clínicas', level=3)
p = doc.add_paragraph()
p.add_run('Dos tablas:').font.size = Pt(11)
p2 = doc.add_paragraph()
p2.add_run('• Distribución del IMC Poblacional: ').bold = True
p2.add_run('Bajo peso · Normal · Sobrepeso · Obesidad (con porcentajes y barra de progreso)')
for r in [p, p2]:
    for run in r.runs:
        run.font.size = Pt(11)

p3 = doc.add_paragraph()
p3.add_run('• Comorbilidades por Sexo: ')
p3.runs[0].bold = True
p3.add_run('tabla cruzada de cada patología/condición desglosada por Masculino y Femenino')
for run in p3.runs:
    run.font.size = Pt(11)

p6 = doc.add_paragraph()
p6.add_run('(Captura de pantalla: tabla de distribución IMC y tabla de comorbilidades por sexo)').italic = True
p6.runs[0].font.size = Pt(10)
p6.runs[0].font.color.rgb = RGBColor(0x66, 0x66, 0x66)

doc.save(r'C:\Users\Bleidys Larios\CLIA\Seccion5_Analitica.docx')
print("OK")
