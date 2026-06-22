from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT

doc = Document()
s = doc.styles['Normal']
s.font.name = 'Calibri'
s.font.size = Pt(11)

doc.add_heading('5. Analítica Descriptiva', level=2)

p = doc.add_paragraph()
p.add_run("La sección de Analítica Descriptiva proporciona un análisis estadístico detallado de los datos clínicos cargados. Se accede desde el sidebar menú ").font.size = Pt(11)
r = p.add_run('Analítica')
r.bold = True
r.font.size = Pt(11)
p.add_run('.').font.size = Pt(11)

doc.add_paragraph()

# 5.1 Prevalencia
doc.add_heading('5.1 Indicadores de Prevalencia Patológica', level=3)
p = doc.add_paragraph()
p.add_run("Se muestran ").font.size = Pt(11)
r = p.add_run('7 tarjetas')
r.bold = True
r.font.size = Pt(11)
p.add_run(' con cantidad y porcentaje de pacientes por condición. Cada tarjeta es cliqueable y abre un modal con la lista filtrada.').font.size = Pt(11)

tabla1 = doc.add_table(rows=8, cols=3)
tabla1.style = 'Light Grid Accent 1'
tabla1.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, (a, b, c) in enumerate([
    ('Indicador', 'Color', 'Icono'),
    ('Hipertensión Arterial', 'danger (rojo)', 'heart-pulse'),
    ('Diabetes Mellitus', 'warning (amarillo)', 'droplet'),
    ('Tabaquismo Activo', 'secondary (gris)', 'smoking'),
    ('Obesidad', 'dark (negro)', 'weight-scale'),
    ('Antecedentes Familiares', 'info (azul)', 'dna'),
    ('Consumo de Alcohol', 'primary (azul)', 'wine-bottle'),
    ('Saturación Baja (<85%)', 'danger (rojo)', 'lungs'),
]):
    for j, txt in enumerate([a, b, c]):
        run = tabla1.cell(i, j).paragraphs[0].add_run(txt)
        run.font.size = Pt(10)
        if i == 0:
            run.bold = True

p = doc.add_paragraph()
r = p.add_run('(Captura de pantalla: tarjetas de prevalencia patológica con barras de progreso)')
r.italic = True
r.font.size = Pt(10)
r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

doc.add_paragraph()

# 5.2 Alertas
doc.add_heading('5.2 Alertas Clínicas', level=3)
p = doc.add_paragraph()
p.add_run('Tres alertas automáticas basadas en umbrales críticos:').font.size = Pt(11)

tabla2 = doc.add_table(rows=4, cols=3)
tabla2.style = 'Light Grid Accent 1'
tabla2.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, (a, b, c) in enumerate([
    ('Alerta', 'Umbral', 'Color'),
    ('Presión Sistólica Alta', '> 180 mmHg', 'rojo'),
    ('Glucosa Alta', '> 300 mg/dL', 'naranja'),
    ('Saturación de Oxígeno Baja', '< 85%', 'azul'),
]):
    for j, txt in enumerate([a, b, c]):
        run = tabla2.cell(i, j).paragraphs[0].add_run(txt)
        run.font.size = Pt(10)
        if i == 0:
            run.bold = True

p = doc.add_paragraph()
r = p.add_run('(Captura de pantalla: tarjetas de alertas clínicas con conteo de pacientes)')
r.italic = True
r.font.size = Pt(10)
r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

doc.add_paragraph()

# 5.3 Segmentacion
doc.add_heading('5.3 Segmentación', level=3)

p = doc.add_paragraph()
r = p.add_run('Por Nivel de Riesgo: ')
r.bold = True
r.font.size = Pt(11)
p.add_run('Tabla con Bajo, Medio, Alto y Crítico, mostrando cantidad, porcentaje y barra de progreso.').font.size = Pt(11)

p = doc.add_paragraph()
r = p.add_run('Por Grupo de Edad: ')
r.bold = True
r.font.size = Pt(11)
p.add_run("Tabla con los grupos etarios '<30', '30-49', '50-69', '70+', cantidad, porcentaje y barra de progreso.").font.size = Pt(11)

p = doc.add_paragraph()
r = p.add_run('(Captura de pantalla: ambas tablas de segmentación lado a lado)')
r.italic = True
r.font.size = Pt(10)
r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

doc.add_paragraph()

# 5.4 Matriz
doc.add_heading('5.4 Matriz Descriptiva de Variables Biomédicas', level=3)
p = doc.add_paragraph()
p.add_run('Tabla con estadísticos descriptivos de las variables continuas: count, media, mediana, moda, desviación estándar, mínimo, máximo, P25 y P75.').font.size = Pt(11)

p = doc.add_paragraph()
r = p.add_run('(Captura de pantalla: matriz descriptiva completa con todas las variables)')
r.italic = True
r.font.size = Pt(10)
r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

doc.add_paragraph()

# 5.5 Distribuciones
doc.add_heading('5.5 Distribuciones Clínicas', level=3)

p = doc.add_paragraph()
r = p.add_run('Distribución del IMC Poblacional: ')
r.bold = True
r.font.size = Pt(11)
p.add_run('Clasificación del IMC (Bajo peso, Normal, Sobrepeso, Obesidad) con cantidad de pacientes y porcentajes.').font.size = Pt(11)

p = doc.add_paragraph()
r = p.add_run('Comorbilidades por Sexo: ')
r.bold = True
r.font.size = Pt(11)
p.add_run('Desglose de condiciones patológicas por Masculino y Femenino, con cantidad y porcentaje.').font.size = Pt(11)

p = doc.add_paragraph()
r = p.add_run('(Captura de pantalla: distribución IMC y comorbilidades por sexo lado a lado)')
r.italic = True
r.font.size = Pt(10)
r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

doc.save(r'C:\Users\Bleidys Larios\CLIA\Seccion5_Analitica.docx')
print("OK")
